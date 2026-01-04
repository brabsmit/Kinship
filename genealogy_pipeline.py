import re
import json
import os
import requests
import time
import hashlib
from docx import Document
from collections import defaultdict
import dateparser
import dateparser.search
from google import genai
from google.genai import types
from dotenv import load_dotenv

class ShipEnrichmentService:
    def __init__(self):
        print("Initializing GenAI API...")
        load_dotenv()
        self.cache_file = "./kinship-app/src/ship_cache.json"
        self.cache = self.load_cache()
        self.cache_updated = False
        self.api_key = os.environ.get("VITE_GEMINI_API_KEY")
        if self.api_key == None:
            self.api_key = os.getenv("VITE_GEMINI_API_KEY")
        self.client = None
        if self.api_key:
            try:
                self.client = genai.Client(api_key=self.api_key)
                print("successfully initialized GenAI API")
            except Exception as e:
                print(f"Warning: Failed to initialize Google GenAI client: {e}")
        else:
            print("Failed to load API key")

    def load_cache(self):
        if os.path.exists(self.cache_file):
            with open(self.cache_file, "r") as f:
                try:
                    return json.load(f)
                except json.JSONDecodeError:
                    return {}
        return {}

    def save_cache(self):
        if self.cache_updated:
            with open(self.cache_file, "w") as f:
                json.dump(self.cache, f, indent=4)
            print("Ship cache saved.")

    def enrich_ship(self, ship_name):
        if not ship_name:
            return None

        # Check Cache
        if ship_name in self.cache:
            return self.cache[ship_name]

        if not self.client:
            return None

        print(f"   [ShipInfo] Querying AI for '{ship_name}'...")
        try:
            prompt = f"""
            Provide historical specifications for the ship named '{ship_name}' (likely 17th-19th century context).
            Return ONLY a JSON object with the following keys. If specific data is unknown, use "Unknown".

            Keys:
            - year_built (string)
            - location_built (string)
            - deck_length (string)
            - beam (string)
            - gross_tonnage (string)
            - masts (string or int)
            - owner (string)
            - description (short string, max 20 words)
            """

            response = self.client.models.generate_content(
                model="gemini-2.0-flash",
                contents=prompt,
                config=types.GenerateContentConfig(
                    response_mime_type="application/json"
                )
            )

            if response.text:
                data = json.loads(response.text)
                # Handle potential list response
                if isinstance(data, list) and len(data) > 0:
                    data = data[0]

                self.cache[ship_name] = data
                self.cache_updated = True
                # Polite delay
                time.sleep(1)
                return data

        except Exception as e:
            print(f"   [ShipInfo] Error enriching '{ship_name}': {e}")

        # Cache miss as None if failed, or maybe we want to retry later?
        # For now, let's not cache failures permanently, or maybe cache a 'checked' state.
        # But for simplicity, we just return None.
        return None

class GeocodingService:
    def __init__(self):
        self.cache_file = "./kinship-app/src/geocoding_cache.json"
        self.cache = self.load_cache()
        self.cache_updated = False
        self.api_enabled = True

        self.HARDCODED_LOCATIONS = {
            "Hartford, CT": [41.7658, -72.6734],
            "Manhattan, NY": [40.7831, -73.9712],
            "New York, NY": [40.7128, -74.0060],
            "New York City": [40.7128, -74.0060],
            "New York": [40.7128, -74.0060],
            "Waterbury, CT": [41.5582, -73.0515],
            "New Haven, CT": [41.3083, -72.9279],
            "Englewood, NJ": [40.8929, -73.9726],
            "Oyster Bay, Long Island, NY": [40.8653, -73.5324],
            "Oyster Bay, NY": [40.8653, -73.5324],
            "Newburgh, NY": [41.5032, -74.0104],
            "Boston, MA": [42.3601, -71.0589],
            "Norwich, CT": [41.5243, -72.0759],
            "Simsbury, CT": [41.8759, -72.8012],
            "Simsbury CT": [41.8759, -72.8012],
            "Middletown, CT": [41.5623, -72.6506],
            "Branford, CT": [41.2795, -72.8151],
            "Stratford, CT": [41.1845, -73.1332],
            "Ipswich, MA": [42.6792, -70.8412],
            "Watertown, MA": [42.3709, -71.1828],
            "Brooklyn, CT": [41.7884, -71.9495],
            "Killingly, CT": [41.8537, -71.8795],
            "Hampton, CT": [41.7834, -72.0531],
            "East Haddam, CT": [41.4587, -72.4626],
            "Windsor, CT": [41.8519, -72.6437],
            "West Hartford, CT": [41.7621, -72.7420],
            "Wethersfield, CT": [41.7145, -72.6579],
            "Westbury, Long Island, NY": [40.7557, -73.5876],
            "Westbury, NY": [40.7557, -73.5876],
            "Colchester, CT": [41.5734, -72.3331],
            "New London, CT": [41.3557, -72.0995],
            "Waltham, MA": [42.3765, -71.2356],
            "Worcester, MA": [42.2626, -71.8023],
            "Salem, MA": [42.5195, -70.8967],
            "Medford, MA": [42.4184, -71.1062],
            "Woburn, MA": [42.4793, -71.1523],
            "England": [52.3555, -1.1743],
            "London, England": [51.5074, -0.1278],
            "Athens, PA": [41.9529, -76.5163],
            "Coeymans, NY": [42.4776, -73.7946],
            "Coxsackie, NY": [42.3601, -73.8068],
            "Shelton, Fairfield County, CT": [41.3165, -73.0932],
            "Milford, CT": [41.2307, -73.0640],
            "Trumbull, CT": [41.2562, -73.1909],
            "Derby, CT": [41.3207, -73.0890],
            "Sunderland, MA": [42.4695, -72.5795],
            "Conway, MA": [42.5106, -72.6976],
            "Bridport, Dorset, England": [50.7337, -2.7563],
            "Dorset, England": [50.7483, -2.3452],
            "Great Limber, Lincolnshire, England": [53.5656, -0.2874],
            "Lincolnshire, England": [53.2285, -0.5478],
            "Somerset, England": [51.0109, -3.1029],
            "Taunton, MA": [41.9001, -71.0898],
            "Essex, England": [51.7670, 0.4664],
            "Cambridge, MA": [42.3736, -71.1097],
            "Hingham, MA": [42.2417, -70.8898],
            "Marshfield, Plymouth Colony, MA": [42.0917, -70.7056],
            "Roxbury, MA": [42.3152, -71.0914],
            "Dedham, MA": [42.2436, -71.1699]
        }

        self.HISTORICAL_LOCATIONS = {
            "New Amsterdam": [40.7128, -74.0060],
            "Massachusetts Bay Colony": [42.3601, -71.0589],
            "Plymouth Colony": [41.9584, -70.6673],
            "New Netherland": [42.6526, -73.7562],
            "Acadia": [44.9667, -64.1333],
            "Prussia": [52.5200, 13.4050],
            "British America": [37.4316, -78.6569],
            "Thirteen Colonies": [39.9526, -75.1652],
            "New France": [46.8139, -71.2080],
            "Saybrook Colony": [41.2934, -72.3898],
            "New Haven Colony": [41.3083, -72.9279],
            "Connecticut Colony": [41.7658, -72.6734],
            "Jamestown": [37.2102, -76.7777],
            "Plymouth": [41.9584, -70.6673],
            "Danzig": [54.3520, 18.6466],
        }

        self.REGION_COORDINATES = {
            "CT": [41.6032, -72.7],
            "Connecticut": [41.6032, -72.7],
            "MA": [42.4072, -71.3824],
            "Massachusetts": [42.4072, -71.3824],
            "NY": [43.0, -75.0],
            "New York": [43.0, -75.0],
            "NJ": [40.0583, -74.4057],
            "New Jersey": [40.0583, -74.4057],
            "PA": [41.2033, -77.1945],
            "Pennsylvania": [41.2033, -77.1945],
            "USA": [39.8283, -98.5795],
            "United States": [39.8283, -98.5795],
            "England": [52.3555, -1.1743],
            "UK": [52.3555, -1.1743],
            "United Kingdom": [52.3555, -1.1743],
            "VA": [37.4316, -78.6569],
            "Virginia": [37.4316, -78.6569],
            "RI": [41.5801, -71.4774],
            "Rhode Island": [41.5801, -71.4774],
            "VT": [44.5588, -72.5778],
            "Vermont": [44.5588, -72.5778],
            "NH": [43.1939, -71.5724],
            "New Hampshire": [43.1939, -71.5724],
            "ME": [45.2538, -69.4455],
            "Maine": [45.2538, -69.4455]
        }

    def load_cache(self):
        if os.path.exists(self.cache_file):
            with open(self.cache_file, "r") as f:
                try:
                    return json.load(f)
                except json.JSONDecodeError:
                    return {}
        return {}

    def save_cache(self):
        if self.cache_updated:
            with open(self.cache_file, "w") as f:
                json.dump(self.cache, f, indent=4)
            print("Geocoding cache saved.")

    def geocode(self, location_name):
        if not location_name or location_name == "Unknown":
            return None

        # Tier 1: Hardcoded
        if location_name in self.HARDCODED_LOCATIONS:
            return {"lat": self.HARDCODED_LOCATIONS[location_name][0], "lng": self.HARDCODED_LOCATIONS[location_name][1], "tier": 1}

        # Tier 2: Historical
        if location_name in self.HISTORICAL_LOCATIONS:
            return {"lat": self.HISTORICAL_LOCATIONS[location_name][0], "lng": self.HISTORICAL_LOCATIONS[location_name][1], "tier": 2}

        # Tier 3: Regions (Exact match)
        if location_name in self.REGION_COORDINATES:
            return {"lat": self.REGION_COORDINATES[location_name][0], "lng": self.REGION_COORDINATES[location_name][1], "tier": 3}

        # Check Cache
        if location_name in self.cache:
            return self.cache[location_name]

        # Tier 4: API
        if not self.api_enabled:
            return None

        try:
            return None
            print(f"   [Geocoding] Querying API for '{location_name}'...")
            time.sleep(1.1)
            url = "https://nominatim.openstreetmap.org/search"
            params = {
                "q": location_name,
                "format": "json",
                "limit": 1
            }
            headers = {"User-Agent": "GenealogyApp/1.0 (contact@example.com)"}
            response = requests.get(url, params=params, headers=headers, timeout=2)
            data = response.json()

            if data:
                lat = float(data[0]["lat"])
                lng = float(data[0]["lon"])
                result = {"lat": lat, "lng": lng, "tier": 4}
                self.cache[location_name] = result
                self.cache_updated = True
                return result
            else:
                self.cache[location_name] = None
                self.cache_updated = True
                return None
        except Exception as e:
            return None

class GenealogyTextPipeline:
    def __init__(self):
        self.family_data = []
        self.image_cache = self.load_cache()
        self.cache_updated = False

    def _extract_voyages(self, text):
        voyages = []
        if not text: return voyages, text

        # Pattern 1: Explicit Tag
        # [Ship: Name | Type: X | Year: Y | Departure: A | Arrival: B]
        tag_pattern = r'\[Ship:\s*([^\]]+)\]'

        def replace_tag(match):
            content = match.group(1)
            parts = [p.strip() for p in content.split('|')]

            voyage = {
                "ship_name": parts[0],
                "type": "Unknown",
                "year": "Unknown",
                "departure": "Unknown",
                "arrival": "Unknown",
                "class": "Passenger"
            }

            for part in parts[1:]:
                if ":" in part:
                    k, v = part.split(":", 1)
                    k = k.strip().lower()
                    v = v.strip()
                    if k == "type": voyage["type"] = v
                    elif k == "year": voyage["year"] = v
                    elif k == "departure": voyage["departure"] = v
                    elif k == "arrival": voyage["arrival"] = v
                    elif k == "class": voyage["class"] = v

            voyages.append(voyage)
            return "" # Remove tag

        new_text = re.sub(tag_pattern, replace_tag, text, flags=re.IGNORECASE)

        # Pattern 2: Natural Language (Improved)
        # Strategy: Find "arrived/sailed/came ... on/aboard ... [ShipName]"

        # 2a. Quoted Ship Name (Strongest, handles "Hector", "Mayflower")
        quote_pattern = r'(?:arrived|sailed|came|passage|travelled)\b[^.;]*?\b(?:on|aboard)\b(?:\s+(?:the|a)\b)?(?:\s*ship)?\s*[“"\'‘]([^”"\'’]+)[”"\'’]'

        for m in re.finditer(quote_pattern, new_text, re.IGNORECASE):
            ship_name = m.group(1).strip()
            if not any(v['ship_name'] == ship_name for v in voyages):
                 voyages.append({
                    "ship_name": ship_name,
                    "type": "Unknown",
                    "year": "Unknown",
                    "departure": "Unknown",
                    "arrival": "Unknown",
                    "class": "Passenger"
                })

        # 2b. Unquoted Capitalized Ship Name (Context aware)
        # Regex Explanation:
        # (?i:...) Group: Case-insensitive context (verbs, prepositions)
        # ([A-Z]...): Case-sensitive Ship Name (must be capitalized)
        prefix = r'(?i:(?:arrived|sailed|came|passage|travelled)\b[^.;]*?\b(?:on|aboard)\b(?:\s+(?:the|a)\b)?(?:\s*ship)?\s+)'
        name_part = r'([A-Z][a-z]+(?:(?:\s+(?:and|&)\s+|\s+)[A-Z][a-z]+)*)'
        cap_pattern = prefix + name_part

        for m in re.finditer(cap_pattern, new_text):
            ship_name = m.group(1).strip()

            # Validation: Ignore common capitalized words that might follow "on"
            ignored = ["May", "June", "July", "August", "September", "October", "November", "December", "Monday", "Sunday", "Christmas", "Easter", "Board"]
            if ship_name in ignored: continue

            if not any(v['ship_name'] == ship_name for v in voyages):
                 voyages.append({
                    "ship_name": ship_name,
                    "type": "Unknown",
                    "year": "Unknown",
                    "departure": "Unknown",
                    "arrival": "Unknown",
                    "class": "Passenger"
                })

        # Cleanup extra spaces
        new_text = re.sub(r'\s{2,}', ' ', new_text).strip()
        return voyages, new_text

    def load_cache(self):
        cache_file = "./kinship-app/src/wikimedia_cache.json"
        if os.path.exists(cache_file):
            with open(cache_file, "r") as f:
                try:
                    return json.load(f)
                except json.JSONDecodeError:
                    return {}
        return {}

    def save_cache(self):
        if self.cache_updated:
            with open("./kinship-app/src/wikimedia_cache.json", "w") as f:
                json.dump(self.image_cache, f, indent=4)
            print("Wikimedia cache saved.")

    def fetch_wikimedia_image(self, location, year):
        if not location or location.lower() == "unknown":
            return None

        # Determine Century
        century = ""
        if year:
            c_val = (year // 100) + 1
            century = f"{c_val}th century"
        else:
            century = "historical"

        # Create a cache key
        cache_key = f"{location}|{century}"
        if cache_key in self.image_cache:
            return self.image_cache[cache_key]

        # Construct Search Queries
        queries = [
            f"{location} {century} map",
            f"{location} historical"
        ]

        api_url = "https://commons.wikimedia.org/w/api.php"

        for q in queries:
            # Check if we need to save cache incrementally
            if self.cache_updated and len(self.image_cache) % 5 == 0:
                self.save_cache()
                self.cache_updated = False

            params = {
                "action": "query",
                "generator": "search",
                "gsrnamespace": "6", # File namespace
                "gsrsearch": q,
                "gsrlimit": "1",
                "prop": "imageinfo",
                "iiprop": "url|extmetadata",
                "iiurlwidth": "1024",
                "format": "json"
            }

            try:
                # Be polite
                time.sleep(0.1)
                response = requests.get(api_url, params=params, headers={"User-Agent": "GenealogyApp/1.0 (contact@example.com)"}, timeout=5)
                data = response.json()

                if "query" in data and "pages" in data["query"]:
                    # Get first result
                    page_id = list(data["query"]["pages"].keys())[0]
                    page = data["query"]["pages"][page_id]

                    if "imageinfo" in page:
                        info = page["imageinfo"][0]
                        thumb_url = info.get("thumburl", info.get("url"))

                        metadata = info.get("extmetadata", {})
                        description = metadata.get("ImageDescription", {}).get("value", q)
                        # Clean HTML from description if needed, or keep it simple
                        # Simple regex to strip HTML tags
                        description = re.sub(r'<[^>]+>', '', description)[:150] + "..."

                        result = {
                            "src": thumb_url,
                            "alt": f"Historical image of {location}",
                            "caption": description,
                            "style": { "filter": "sepia(20%) contrast(110%)" } # Default vintage style
                        }

                        self.image_cache[cache_key] = result
                        self.cache_updated = True
                        print(f"   [Wikimedia] Found image for '{q}'")
                        return result

            except Exception as e:
                print(f"   [Wikimedia] Error fetching for '{q}': {e}")
                continue

        # Cache miss (None) to avoid re-searching
        self.image_cache[cache_key] = None
        self.cache_updated = True
        return None

    def _normalize_date(self, raw_date_string):
        """
        Parses a raw date string and returns a best-guess integer year.
        Returns None if no valid year is found.
        """
        if not raw_date_string:
            return None

        # Clean up the string
        s = raw_date_string.strip().lower()
        if s in ["unknown", "?", "uncertain", "possibly", ""]:
            return None

        # Extract the first 4-digit year candidate to work with
        # (1000-2999).
        # We capture the group to ensure we get the year digits.
        year_match = re.search(r'\b(1[0-9]{3}|20[0-2][0-9])', s)

        # 0. Handle "century" logic if no specific year found
        if not year_match:
            # "18th century" -> 1700
            # "17th century" -> 1600
            century_match = re.search(r'\b(\d{2})(?:th|nd|st|rd)\s+century\b', s)
            if century_match:
                try:
                    c_val = int(century_match.group(1))
                    return (c_val - 1) * 100
                except:
                    pass

            # Fallback to dateparser if no 4-digit year found (e.g. "Apr 12, '80")
            return self._normalize_date_fallback(raw_date_string)

        year_val = int(year_match.group(1))
        start_index = year_match.start()

        # Look at the text *before* the year for modifiers
        pre_text = s[:start_index]

        # 1. Handle "before" / "bef" / "by"
        # Logic: if "bef", "before", or "by" appears in the text preceding the year, return year - 1
        # e.g., "bef 1800" -> 1799
        if re.search(r'\b(bef\.?|before|by)\b', pre_text):
            return year_val - 1

        # 2. Handle "after" / "aft"
        # Logic: if "aft" or "after" appears, return year + 1
        # e.g., "aft 1750" -> 1751
        if re.search(r'\b(aft\.?|after)\b', pre_text):
            return year_val + 1

        # 3. Handle "between"
        # Logic: "between 1770 and 1780" -> 1770 (Start date).
        # This falls through to the default return of year_val because regex found the first year.
        # But we verify no "bef" / "aft" modifiers confuse it.
        # "between" in pre_text -> simply return year_val.
        if re.search(r'\bbetween\b', pre_text):
            return year_val

        # 4. Handle dual dating like "1774/5" or ranges "1774-1778"
        # The regex picks the first year found, which is standard genealogical practice for sorting (start date).
        # For "1774/5", it extracts "1774".

        # 5. Handle "living in" or "fl."
        # If "living in 1774", we return 1774 as the best anchor.
        if re.search(r'\b(living in|fl\.?)\b', pre_text):
            return year_val

        # 6. Handle "circa" / "c." / "about" / "abt"
        # If "c. 1774", we extract 1774.
        if re.search(r'\b(c\.?|ca\.?|circa|about|abt\.?)\b', pre_text):
            return year_val

        return year_val

    def _normalize_date_fallback(self, raw_date_string):
        """
        Uses dateparser to attempt to find a year if regex failed.
        """
        try:
            dt = dateparser.parse(raw_date_string)
            if dt:
                return dt.year
        except:
            pass
        return None

    def split_date_location(self, text):
        if not text or text.lower() == "unknown":
            return "Unknown", "Unknown"

        # 1. " in " separator (Strongest)
        # Handles: "May 1, 1850 in Hartford" -> "May 1, 1850", "Hartford"
        in_sep = re.search(r"\s+in\s+", text, re.IGNORECASE)
        if in_sep:
            parts = re.split(r"\s+in\s+", text, flags=re.IGNORECASE, maxsplit=1)
            date_candidate = parts[0].strip()
            loc_candidate = parts[1].strip()

            # Correction: If location is just a year (e.g. "Disappeared in 1744"),
            # then the whole thing is likely a date statement, or at least the year belongs to the date.
            # We treat the whole string as the date, and location as Unknown.
            if re.match(r'^\d{4}$', loc_candidate):
                 return text.strip(), "Unknown"

            # Clean up date_candidate: Remove standard event labels if present
            # (e.g. "Born: April 12, 1880" -> "April 12, 1880")
            date_candidate = re.sub(r'^(Born|Died|Buried|Baptized|Married):?\s*', '', date_candidate, flags=re.IGNORECASE)

            # Cleanup for "possibly" or garbage
            if date_candidate.strip().lower() in ["possibly", "unknown", "?", ""]:
                 date_candidate = "Unknown"

            return date_candidate, loc_candidate

        # 2. Fallback: Use dateparser to handle messy formats
        # e.g., "Springfield, 1880", "Born 1880", "1880 New York"
        try:
            dates = dateparser.search.search_dates(text, languages=['en'])
        except Exception:
            dates = None

        if dates:
            # Found one or more dates.
            # We want to identify the span of the date substring(s) and treat the rest as location.

            # Start of the first date match
            first_date_str = dates[0][0]
            start_idx = text.find(first_date_str)

            # End of the last date match
            last_date_str = dates[-1][0]
            # Use rfind to be safe in case of repeated strings, but restrict search after start_idx?
            # Actually, standard find for last_date_str is fine if we assume chronological parsing order,
            # but safer to find last occurrence if dateparser returns them in order found in text.
            # dateparser search_dates returns list in order of appearance.
            end_idx = text.rfind(last_date_str) + len(last_date_str)

            # Expand left to capture modifiers that dateparser might miss (c., Before, etc.)
            pre_text = text[:start_idx]
            mod_pattern = r'(?i)\b(?:c\.?|ca\.?|circa|about|abt\.?|before|bef\.?|by|after|aft\.?|bet\.?|between|living\s+in|fl\.?)\s*$'
            mod_match = re.search(mod_pattern, pre_text)

            if mod_match:
                start_idx = mod_match.start()

            date_part = text[start_idx:end_idx].strip()

            # Clean "in" / "at" from end of date_part if dateparser captured it
            # e.g. "April 12, 1880 in"
            date_part = re.sub(r'\s+(in|at|on)$', '', date_part, flags=re.IGNORECASE)

            # Extract Location (everything else)
            prefix = text[:start_idx].strip()
            suffix = text[end_idx:].strip()

            # Clean up suffix
            suffix = re.sub(r'^(in|at|on)\b\s*', '', suffix, flags=re.IGNORECASE)

            # Clean up prefix: Remove standard event labels if they were part of the string
            # (e.g. "Born: April 12" -> Prefix "Born: ")
            prefix = re.sub(r'^(Born|Died|Buried|Baptized|Married):?\s*', '', prefix, flags=re.IGNORECASE)

            # Combine prefix and suffix
            loc_parts = []
            if prefix.strip(",; "): loc_parts.append(prefix.strip(",; "))
            if suffix.strip(",; "): loc_parts.append(suffix.strip(",; "))

            location = ", ".join(loc_parts)
            if not location:
                location = "Unknown"

            return date_part, location

        # 3. No date found by dateparser
        # Heuristics for "No Year" or fallbacks

        lower_text = text.lower()

        # Specific cleanup for "Unknown date, assume..."
        if "unknown date" in lower_text or "date unknown" in lower_text:
             # Strip that part out
             cleaned = re.sub(r'\b(unknown date|date unknown)\b', '', text, flags=re.IGNORECASE).strip(" ,;.")
             # Assume the rest is location if it has content
             if cleaned:
                  return "Unknown", cleaned

        # Explicitly suppress non-location garbage words if they appear alone
        garbage_words = ["possibly", "unknown", "uncertain", "?", ""]
        if lower_text.strip(" .,;") in garbage_words:
             return "Unknown", "Unknown"

        # General Keyword Check
        keywords = ["unknown", "?", "disappeared", "uncertain", "infant"]
        has_keyword = any(k in lower_text for k in keywords)

        has_digits = any(char.isdigit() for char in text)

        if has_digits:
             # Likely a date that dateparser missed?
             return text.strip(), "Unknown"

        if has_keyword:
             # It has a keyword like "?". Is it "England?" or "?"
             # If it has valid words (length > 3), treat as location?
             # "England?" -> Location: England?
             # "?" -> Date: ?, Loc: Unknown

             clean_text = text.strip(" ?.")
             if len(clean_text) > 3:
                  # Likely a location with a question mark
                  return "Unknown", text.strip()
             else:
                  # Just "?" or "Unk"
                  return text.strip(), "Unknown"

        # Treat as Location if no digits
        return "Unknown", text.strip()

    def _parse_location_hierarchy(self, location_string):
        """
        Parses a location string into structured components (City, State, Country).
        Uses simple comma-splitting heuristics.
        """
        if not location_string or location_string.lower() == "unknown":
            return None

        parts = [p.strip() for p in location_string.split(',')]

        hierarchy = {
            "raw": location_string,
            "city": None,
            "county": None,
            "state": None,
            "country": None
        }

        # Common US State Abbreviations/Names appearing in this dataset
        us_states = ["CT", "MA", "NY", "NJ", "PA", "VA", "RI", "NH", "VT", "ME", "DE", "MD", "SC", "NC", "GA", "OH", "IL", "MI"]

        if len(parts) == 1:
            # Just a city or country?
            val = parts[0]
            if val in ["USA", "England", "UK", "Canada"]:
                hierarchy["country"] = val
            elif val in us_states:
                hierarchy["state"] = val
                hierarchy["country"] = "USA"
            else:
                hierarchy["city"] = val

        elif len(parts) >= 2:
            last = parts[-1]

            # Check if last part is a US State
            if last.upper() in us_states or last in ["Connecticut", "Massachusetts", "New York", "New Jersey", "Pennsylvania"]:
                hierarchy["country"] = "USA"
                hierarchy["state"] = last
                hierarchy["city"] = parts[0]
                if len(parts) > 2:
                    hierarchy["county"] = parts[1] # Middle part often county

            elif last in ["USA", "United States"]:
                hierarchy["country"] = "USA"
                hierarchy["state"] = parts[-2] if len(parts) > 1 else None
                hierarchy["city"] = parts[0]

            elif last in ["England", "UK", "Great Britain"]:
                hierarchy["country"] = "UK"
                hierarchy["city"] = parts[0]
                if len(parts) > 2:
                    hierarchy["county"] = parts[1]

            else:
                # Generic fallback: City, Region
                hierarchy["city"] = parts[0]
                hierarchy["state"] = parts[-1]

        return hierarchy

    def _extract_location_note(self, location_text):
        """
        Separates a trailing parenthetical note from the location string.
        Returns (clean_location, note).
        """
        if not location_text or location_text == "Unknown":
            return location_text, None

        # Regex to find text ending with (note)
        # We capture the main part (lazy) and the content inside the *last* set of parens.
        # This handles cases like "Location (Note)" -> "Location", "Note"
        match = re.search(r'^(.*?)\s*\(([^)]+)\)$', location_text)
        if match:
            clean_loc = match.group(1).strip()
            note = match.group(2).strip()
            # If the resulting location is empty, it means the whole string was in parens
            # e.g. "(Unknown)" -> loc="", note="Unknown".
            # In this case, we probably shouldn't split it, or treat it as just note?
            # For now, if clean_loc is empty, we return the original string to avoid data loss/empty location.
            if not clean_loc:
                return location_text, None

            return clean_loc, note

        return location_text, None

    def split_sentences(self, text):
        if not text:
            return []
        text = re.sub(r"^NOTES:\s*", "", text, flags=re.IGNORECASE)
        parts = re.split(r'[.;]\s+', text)
        return [p.strip() for p in parts if p.strip()]

    def extract_year(self, text):
        matches = list(re.finditer(r'\b(16|17|18|19)\d{2}\b', text))
        if matches:
            return int(matches[0].group(0))
        return None

    def extract_relative_date(self, text, context_year):
        if not context_year:
            return None
        match = re.search(r'\b(one|two|three|four|five|six|seven|eight|nine|ten|twenty|thirty|forty|fifty)\s+years?\s+(later|after)', text, re.IGNORECASE)
        if match:
            number_map = {
                "one": 1, "two": 2, "three": 3, "four": 4, "five": 5,
                "six": 6, "seven": 7, "eight": 8, "nine": 9, "ten": 10,
                "twenty": 20, "thirty": 30, "forty": 40, "fifty": 50
            }
            num = number_map.get(match.group(1).lower())
            if num:
                return context_year + num
        return None

    def extract_location(self, text):
        loc_regex = r'(?:\b(?:in|at|to|from)\s+)([A-Z][a-zA-Z]+(?:[\s,]+[A-Z][a-zA-Z]+)*)'
        m = re.search(loc_regex, text)
        if m:
            candidate = m.group(1)
            candidate = re.sub(r'\s+in\s+\d{4}.*', '', candidate)
            ignored = ["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December", "Harvard", "Yale", "College", "University", "War", "Church"]
            if candidate in ignored:
                return None
            return candidate
        return None

    def extract_events_from_text(self, text):
        sentences = self.split_sentences(text)
        events = []
        current_context_year = None

        for sent in sentences:
            year = self.extract_year(sent)
            if year:
                current_context_year = year
            else:
                year = self.extract_relative_date(sent, current_context_year)
                if year:
                    current_context_year = year

            if year:
                loc = self.extract_location(sent)
                events.append({
                    "year": year,
                    "label": sent,
                    "location": loc or "Unknown",
                    "type": "personal"
                })
        return events

    def _extract_explicit_associates(self, text):
        associates = []
        # Pattern: [Role: Name (Context)]
        # Supported Roles: Associate, Witness, Neighbor, Partner, Friend
        pattern = r'\[(Associate|Witness|Neighbor|Partner|Friend)\s*:\s*([^\]]+)\]'

        matches = re.finditer(pattern, text, re.IGNORECASE)
        for m in matches:
            role = m.group(1).title()
            content = m.group(2).strip()

            # Split name and optional paren context
            name_match = re.search(r'^(.*?)\s*\(([^)]+)\)$', content)
            if name_match:
                name = name_match.group(1).strip()
                context = name_match.group(2).strip()
            else:
                name = content
                context = None

            associates.append({
                "name": name,
                "role": role,
                "context": context
            })
        return associates

    def parse_document(self, docx_path, lineage_label):
        print(f"--- Scanning Narrative Document ({docx_path}) ---")
        try:
            doc = Document(docx_path)
        except Exception as e:
            print(f"CRITICAL ERROR: Could not load Word doc. {e}")
            return

        total_paragraphs = len(doc.paragraphs)
        print(f"Total paragraphs in document: {total_paragraphs}")

        current_profiles = []
        seen_ids = set()
        current_generation = "Uncategorized"
        
        # Regex Patterns
        id_pattern = re.compile(r"\{(\d+(\.\d+)*)\}")
        born_pattern = re.compile(r"Born:\s*(.*)", re.IGNORECASE)
        died_pattern = re.compile(r"Died:\s*(.*)", re.IGNORECASE)
        children_pattern = re.compile(r"Children:\s*(.*)", re.IGNORECASE)
        notes_start_pattern = re.compile(r"NOTES:\s*(.*)", re.IGNORECASE)
        source_tag_pattern = re.compile(r"\[source:\s*(.*?)\]", re.IGNORECASE)
        
        gen_header_pattern = re.compile(r"^(GENERATION\s+[IVXLCDM]+.*)", re.IGNORECASE)

        for index, para in enumerate(doc.paragraphs):
            if index % 1000 == 0:
                print(f"Processing paragraph {index}/{total_paragraphs}")

            text = para.text.strip()
            if not text:
                continue

            gen_match = gen_header_pattern.match(text)
            if gen_match:
                current_generation = gen_match.group(1).strip()
                print(f"   > Detected Section: {current_generation}")
                if current_profiles:
                    self.family_data.extend(current_profiles)
                    current_profiles = []
                continue

            matches = list(id_pattern.finditer(text))

            # Guard: If the line is actually a Note or Vital Stat line that happens to reference an ID,
            # ignore it as a profile header.
            # Also exclude cross-reference lines like "See Name..." or relationship pointers "Father of..."
            see_pattern = re.compile(r"^See\s+", re.IGNORECASE)
            rel_pointer_pattern = re.compile(r"^(Father|Mother|Parent|Maternal|Paternal)\s+(Grand)?(father|mother|parent|of)\s+", re.IGNORECASE)

            is_metadata_line = (
                born_pattern.match(text) or
                died_pattern.match(text) or
                children_pattern.match(text) or
                notes_start_pattern.match(text) or
                see_pattern.match(text) or
                rel_pointer_pattern.match(text)
            )

            if matches and not is_metadata_line:
                if current_profiles:
                    self.family_data.extend(current_profiles)
                    current_profiles = []

                # Handle First ID (always create)
                first_match = matches[0]

                # Name is everything before the first ID
                raw_name = text[:first_match.start()].strip()
                clean_name = re.sub(r"\[source:.*?\]", "", raw_name).strip()
                
                source_match = source_tag_pattern.search(text)
                source_id = source_match.group(1) if source_match else "Unknown"

                # Check First ID
                uid = first_match.group(1)
                if uid not in seen_ids:
                    seen_ids.add(uid)
                    current_profiles.append({
                        "id": uid,
                        "name": clean_name,
                        "lineage": lineage_label,
                        "generation": current_generation,
                        "vital_stats": {
                            "born_date": "Unknown",
                            "born_location": "Unknown",
                            "died_date": "Unknown",
                            "died_location": "Unknown"
                        },
                        "story": {
                            "notes": "",
                            "voyages": [],
                            "life_events": []
                        },
                        "metadata": {
                            "source_id": source_id,
                            "doc_paragraph_index": index + 1
                        }
                    })

                # Check subsequent IDs (aliases)
                # Logic: Consecutive IDs separated by " & ", " / ", " and "
                for i in range(1, len(matches)):
                    prev = matches[i-1]
                    curr = matches[i]

                    # Text between matches
                    between = text[prev.end():curr.start()]

                    if re.match(r"^\s*(&|/|and)\s*$", between, re.IGNORECASE):
                        # It is an alias
                        uid_alias = curr.group(1)
                        if uid_alias not in seen_ids:
                            seen_ids.add(uid_alias)
                            current_profiles.append({
                                "id": uid_alias,
                                "name": clean_name, # Same name
                                "lineage": lineage_label,
                                "generation": current_generation,
                                "vital_stats": {
                                    "born_date": "Unknown",
                                    "born_location": "Unknown",
                                    "died_date": "Unknown",
                                    "died_location": "Unknown"
                                },
                                "story": {
                                    "notes": "",
                                    "voyages": [],
                                    "life_events": []
                                },
                                "metadata": {
                                    "source_id": source_id,
                                    "doc_paragraph_index": index + 1
                                }
                            })
                    else:
                        # Break chain if separator is not alias-like
                        break

                continue

            if current_profiles:
                # Pre-calculate matches for this line
                b_match = born_pattern.search(text)
                d_match = died_pattern.search(text)
                n_match = notes_start_pattern.search(text)
                c_match = children_pattern.search(text)

                for current_profile in current_profiles:
                    if b_match:
                        raw_born = b_match.group(1).strip()
                        b_date, b_loc_raw = self.split_date_location(raw_born)
                        b_loc, b_note = self._extract_location_note(b_loc_raw)

                        current_profile["vital_stats"]["born_date"] = b_date
                        current_profile["vital_stats"]["born_location"] = b_loc
                        if b_note:
                            current_profile["vital_stats"]["born_location_note"] = b_note
                        current_profile["vital_stats"]["born_year_int"] = self._normalize_date(b_date)
                        current_profile["vital_stats"]["born_hierarchy"] = self._parse_location_hierarchy(b_loc)

                    if d_match:
                        raw_died = d_match.group(1).strip()
                        d_date, d_loc_raw = self.split_date_location(raw_died)
                        d_loc, d_note = self._extract_location_note(d_loc_raw)

                        current_profile["vital_stats"]["died_date"] = d_date
                        current_profile["vital_stats"]["died_location"] = d_loc
                        if d_note:
                            current_profile["vital_stats"]["died_location_note"] = d_note
                        current_profile["vital_stats"]["died_year_int"] = self._normalize_date(d_date)
                        current_profile["vital_stats"]["died_hierarchy"] = self._parse_location_hierarchy(d_loc)

                    if n_match:
                        notes_text = n_match.group(1).strip()
                        voyages, cleaned_notes = self._extract_voyages(notes_text)
                        current_profile["story"]["notes"] = cleaned_notes
                        current_profile["story"]["voyages"] = voyages
                        current_profile["story"]["life_events"] = self.extract_events_from_text(cleaned_notes)
                        current_profile["story"]["associates"] = self._extract_explicit_associates(cleaned_notes)

                    if c_match:
                        raw_children = c_match.group(1).strip()
                        # Also check next paragraphs if they look like list items or continuation?
                        # For now, simplistic parsing of the line
                        self._parse_children(raw_children, current_profile, lineage_label, current_generation, index)

        if current_profiles:
            self.family_data.extend(current_profiles)

    def _parse_children(self, text, parent_profile, lineage_label, generation, index):
        """
        Parses the children text block and creates sibling profiles.
        Format variations:
        - "Name (Year); Name (Year)"
        - "Name (Year-Year)"
        - "Name [Spouse] (Year)"
        """
        if not text: return

        # Split by semicolon usually separates distinct children entries
        # If no semicolon, maybe commas? But names have commas (Last, First).
        # Heuristic: split by semicolon first.
        segments = [s.strip() for s in text.split(';') if s.strip()]

        # If no semicolons, it might be a single child or comma separated?
        # Check if there are years in parens to guide splitting?
        # For now, stick to semicolons or newlines (handled by loop above but here text is one line)

        child_count = 0
        for segment in segments:
            # Extract Name and Year
            # Regex: Name (Date)
            # Name might include [Spouse]

            # Find parens with digits (dates)
            date_match = re.search(r'\(([^)]*\d+[^)]*)\)', segment)

            raw_date = "Unknown"
            clean_name = segment

            if date_match:
                raw_date = date_match.group(1)
                # Remove the date part from name
                clean_name = segment.replace(f"({raw_date})", "").strip()

            # Clean name further (remove " and " or "others")
            if "others" in clean_name.lower():
                continue

            clean_name = clean_name.strip()
            if not clean_name: continue

            # Generate a unique ID for this child
            # Format: P<ParentID>_c<Index>
            # Note: Parent ID must exist.
            if not parent_profile or 'id' not in parent_profile:
                continue

            # Check if this child already exists as a main profile?
            # We will handle this in deduplication or linking phase,
            # but we need to create the object first.

            child_id = f"{parent_profile['id']}_c{child_count}"
            child_count += 1

            # Create Child Profile
            child_profile = {
                "id": child_id,
                "name": clean_name,
                "lineage": lineage_label,
                "generation": generation, # Technically next gen down, but close enough for now
                "vital_stats": {
                    "born_date": raw_date, # Often range "1884-1976" or just "1884"
                    "born_location": "Unknown",
                    "died_date": "Unknown", # Extracted later from range
                    "died_location": "Unknown"
                },
                "story": {
                    "notes": f"Child of {parent_profile['name']}. Source text: {segment}",
                    "life_events": []
                },
                "metadata": {
                    "source_id": "Derived",
                    "doc_paragraph_index": index + 1,
                    "is_child_entry": True,
                    "parent_id": parent_profile['id']
                }
            }

            # Handle Date Ranges in born_date
            # e.g. "1884-1976" -> born 1884, died 1976
            if "-" in raw_date or "–" in raw_date: # Handle en-dash too
                split_char = "–" if "–" in raw_date else "-"
                parts = raw_date.split(split_char)
                if len(parts) >= 2:
                    child_profile["vital_stats"]["born_date"] = parts[0].strip()
                    child_profile["vital_stats"]["died_date"] = parts[1].strip()

            # Clean child dates if they are garbage
            if child_profile["vital_stats"]["born_date"].strip() in ["?", ""]:
                 child_profile["vital_stats"]["born_date"] = "Unknown"
            if child_profile["vital_stats"]["died_date"].strip() in ["?", ""]:
                 child_profile["vital_stats"]["died_date"] = "Unknown"

            child_profile["vital_stats"]["born_year_int"] = self._normalize_date(child_profile["vital_stats"]["born_date"])
            child_profile["vital_stats"]["died_year_int"] = self._normalize_date(child_profile["vital_stats"]["died_date"])

            self.family_data.append(child_profile)

        print(f"Successfully extracted {len(self.family_data)} profiles from text.")

    def _has_exclusion_context(self, text, match_start):
        # Look at the 100 chars before the match (increased from 50)
        start_search = max(0, match_start - 100)
        pre_text = text[start_search:match_start].lower()

        exclusions = [
            # Singular relationships
            "mother of", "father of", "sister of", "brother of",
            "wife of", "husband of", "widow of", "son of", "daughter of",
            "child of", "spouse of", "married to", "mother to", "father to",
            "husband was", "father was", "son was", "daughter was", "wife was",
            "husband's", "father's", "wife's", "son's", "daughter's",
            "consort of", "relict of",
            # Plural relationships
            "sons of", "daughters of", "children of", "brothers of", "sisters of",
            # Contextual exclusions for spouse death scenarios
            "husband died", "wife died", "spouse died"
        ]

        for exc in exclusions:
            if exc in pre_text:
                return True
        return False

    def extract_tags(self, profile):
        tags = []
        notes = profile["story"]["notes"]
        born_loc = profile["vital_stats"]["born_location"]
        died_loc = profile["vital_stats"]["died_location"]

        # 1. Immigrant
        # Heuristic: Born in UK/Europe, Died in USA/MA/CT
        # Or keyword "immigrant", "came to america"
        is_immigrant = False
        immigrant_keywords = r'\b(immigrant|emigrated|came to america|arrived in|arrived with)\b'
        for match in re.finditer(immigrant_keywords, notes, re.IGNORECASE):
            if not self._has_exclusion_context(notes, match.start()):
                is_immigrant = True
                break

        # Simple location check (Location based check is usually safe from context errors)
        uk_locations = ["England", "UK", "Britain", "London", "Dorset", "Essex", "Somerset", "Lincolnshire", "Suffolk", "Kent", "Holland", "Netherlands", "Scotland", "Ireland", "Wales"]
        us_locations = ["MA", "CT", "NY", "NJ", "USA", "Massachusetts", "Connecticut", "New York", "Pennsylvania", "New Hampshire", "Rhode Island"]

        born_uk = any(l.lower() in born_loc.lower() for l in uk_locations)
        died_us = any(l.lower() in died_loc.lower() for l in us_locations)

        if born_uk and died_us:
            is_immigrant = True

        if is_immigrant:
            tags.append("Immigrant")

        # 2. Mayflower
        mayflower_keywords = r'\bMayflower\b'
        for match in re.finditer(mayflower_keywords, notes, re.IGNORECASE):
             if not self._has_exclusion_context(notes, match.start()):
                 tags.append("Mayflower")
                 break

        # 3. War Veteran
        # Keywords: War, Revolution, Army, Regiment, Captain, Lieutenant, General, Soldier, Private
        war_keywords = r'\b(served in|soldier|captain|major|lieutenant|general|ensign|private|sergeant|colonel|veteran|war of|revolutionary war|civil war|french and indian war)\b'
        for match in re.finditer(war_keywords, notes, re.IGNORECASE):
             if not self._has_exclusion_context(notes, match.start()):
                 tags.append("War Veteran")
                 break

        # Name check for rank
        name_keywords = r'captain|major|lieutenant|Lt\.|general|ensign|private|sergeant|colonel'
        if re.search(name_keywords, profile["name"], re.IGNORECASE):
             # Basic check to avoid Mrs. Captain
             if not re.search(r'\b(mrs|miss|ms)\b', profile["name"], re.IGNORECASE):
                tags.append("War Veteran")

        # 4. Founder / Settler
        founder_keywords = r'\b(founder|settler|pioneer|first settler|original proprietor)\b'
        for match in re.finditer(founder_keywords, notes, re.IGNORECASE):
             if not self._has_exclusion_context(notes, match.start()):
                 tags.append("Founder")
                 break

        # 5. Salem Witch Trials
        witch_keywords = r'\b(witch|salem trials|accused of witchcraft)\b'
        for match in re.finditer(witch_keywords, notes, re.IGNORECASE):
             if not self._has_exclusion_context(notes, match.start()):
                 tags.append("Salem Witch Trials")
                 break

        # 6. Education
        edu_keywords = r'\b(Harvard|Yale|College|University)\b'
        for match in re.finditer(edu_keywords, notes, re.IGNORECASE):
             if not self._has_exclusion_context(notes, match.start()):
                 tags.append("University Educated")
                 break

        # 7. Quaker
        quaker_keywords = r'\b(Quaker|Friends)\b'
        for match in re.finditer(quaker_keywords, notes, re.IGNORECASE):
             if not self._has_exclusion_context(notes, match.start()):
                 tags.append("Quaker")
                 break

        # 8. Religious Leader
        # Check Name for "Reverend" or "Deacon"
        if re.search(r'\b(Reverend|Deacon)\b|\bRev\.?', profile["name"], re.IGNORECASE):
             tags.append("Religious Leader")

        # 9. Westward Pioneer
        # Born in East Coast, Died in Midwest/West
        east_coast = ["CT", "MA", "NY", "NJ", "PA", "VA", "RI", "VT", "NH", "ME", "DE", "MD", "Connecticut", "Massachusetts", "New York", "New Jersey", "Pennsylvania", "Virginia", "Rhode Island", "Vermont", "New Hampshire", "Maine"]
        westward_states = ["OH", "IL", "MI", "IN", "WI", "MN", "IA", "MO", "KS", "NE", "SD", "ND", "CA", "OR", "WA", "NV", "AZ", "NM", "UT", "CO", "WY", "ID", "MT", "Ohio", "Illinois", "Michigan", "Indiana", "Wisconsin", "Minnesota", "Iowa", "Missouri", "Kansas", "Nebraska", "California", "Oregon", "Washington", "Nevada", "Arizona", "Utah", "Colorado"]

        # Check if born in East Coast and Died in West
        def is_in_region(loc, region_list):
            if not loc or loc == "Unknown": return False
            for r in region_list:
                # Check for full name or abbreviation with boundary
                if len(r) == 2:
                    if re.search(r'\b' + r + r'\b', loc): return True
                else:
                    if r.lower() in loc.lower(): return True
            return False

        if is_in_region(born_loc, east_coast) and is_in_region(died_loc, westward_states):
            tags.append("Westward Pioneer")

        return list(set(tags))

    def link_family_members(self):
        print("--- Linking Family Members ---")
        id_map = {p['id']: p for p in self.family_data}

        for p in self.family_data:
            pid = p['id']
            if 'relations' not in p:
                 p['relations'] = {
                    "parents": [],
                    "children": [],
                    "spouses": []
                }

            # Skip implicit linking for child entries (IDs with '_c')
            # because they don't follow the strict structural hierarchy
            if '.' in pid and '_c' not in pid:
                child_id = pid.rsplit('.', 1)[0]
                if child_id in id_map:
                    if child_id not in p['relations']['children']:
                        p['relations']['children'].append(child_id)

                    child_p = id_map[child_id]
                    if 'relations' not in child_p:
                        child_p['relations'] = {"parents": [], "children": [], "spouses": []}
                    if pid not in child_p['relations']['parents']:
                        child_p['relations']['parents'].append(pid)

            spouse_candidate = None
            if '.' in pid:
                if pid.endswith('.1'):
                    spouse_candidate = pid[:-1] + '2'
                elif pid.endswith('.2'):
                    spouse_candidate = pid[:-1] + '1'
            else:
                try:
                    nid = int(pid)
                    if nid % 2 != 0:
                        spouse_candidate = str(nid + 1)
                    else:
                        spouse_candidate = str(nid - 1)
                except ValueError:
                    pass

            if spouse_candidate and spouse_candidate in id_map:
                if spouse_candidate not in p['relations']['spouses']:
                    p['relations']['spouses'].append(spouse_candidate)

    def _analyze_naming_patterns(self):
        print("--- Analyzing Naming Patterns (The Echo) ---")
        id_map = {p['id']: p for p in self.family_data}
        count = 0

        for p in self.family_data:
            # 1. Parse Name
            # Remove suffixes for clean splitting
            clean_name = re.sub(r',?\s+(Jr\.?|Sr\.?|III|IV|Esq\.?)$', '', p['name'], flags=re.IGNORECASE)
            parts = clean_name.split()

            # Need at least 3 parts: First Middle Last
            if len(parts) < 3:
                continue

            # Middle name is usually the second token
            # If multiple middle names (e.g. William Earl Dodge), we check all "middle" tokens
            # But usually sticking to the first middle token is safe for "Mother's Maiden Name" logic
            # Let's check all middle tokens just in case
            middle_tokens = parts[1:-1]
            my_surname = parts[-1].lower()

            match_found = None

            for token in middle_tokens:
                token_clean = token.strip(".,")
                if len(token_clean) < 3: continue # Skip initials

                # 2. Traverse Ancestors (BFS)
                queue = [(pid, 1) for pid in p.get('relations', {}).get('parents', [])]
                visited = set()

                while queue:
                    curr_id, steps = queue.pop(0)
                    if curr_id in visited: continue
                    visited.add(curr_id)

                    if curr_id not in id_map: continue
                    ancestor = id_map[curr_id]

                    # Extract Ancestor Surnames
                    # 1. Last Name
                    anc_parts = re.sub(r',?\s+(Jr\.?|Sr\.?|III|IV|Esq\.?)$', '', ancestor['name'], flags=re.IGNORECASE).split()
                    anc_surname = anc_parts[-1].strip(".,")

                    # 2. Maiden Name in Parens: "Mary (Greene) Wainwright"
                    maiden_match = re.search(r'\((.*?)\)', ancestor['name'])
                    anc_maiden = maiden_match.group(1).strip() if maiden_match else ""

                    # Check for Match
                    is_match = False
                    matched_surname = ""

                    if token_clean.lower() == anc_surname.lower():
                        is_match = True
                        matched_surname = anc_surname
                    elif anc_maiden and token_clean.lower() == anc_maiden.lower():
                        is_match = True
                        matched_surname = anc_maiden

                    # Filter out matches that are just the profile's own surname (e.g. Father's surname)
                    # We are looking for "Hidden Logic" - so usually a different name.
                    if is_match and matched_surname.lower() != my_surname:

                        # Determine Relation Label
                        rel_label = "Ancestor"
                        if steps == 1: rel_label = "Mother" # If surname is different from mine, it's likely mother
                        elif steps == 2: rel_label = "Grandmother"
                        elif steps == 3: rel_label = "Great-Grandmother"

                        # Store it
                        match_found = {
                            "middle_name": token_clean,
                            "ancestor_name": ancestor['name'],
                            "ancestor_id": ancestor['id'],
                            "relation": rel_label,
                            "surname": matched_surname
                        }
                        break # Stop looking for this token (found nearest ancestor)

                    # Continue traversal
                    if steps < 5 and not match_found:
                         for par_id in ancestor.get('relations', {}).get('parents', []):
                             queue.append((par_id, steps + 1))

                if match_found:
                    break # Stop checking other middle tokens

            if match_found:
                p['story']['naming_echo'] = match_found
                count += 1

        print(f"Found {count} Naming Echo matches.")

    def _get_birth_year(self, profile):
        raw = profile.get("vital_stats", {}).get("born_date", "")
        match = re.search(r'\d{4}', raw)
        if match:
            return int(match.group(0))
        return None

    def _build_name_index(self):
        """
        Builds a comprehensive index of names to profile IDs, handling
        variations like suffixes (Jr., Sr.) and middle names.
        """
        name_index = defaultdict(list)

        for p in self.family_data:
            pid = p['id']
            full_name = p['name']

            # Clean: remove [source], {id}, and trailing punctuation
            clean_name = re.sub(r'\[.*?\]', '', full_name).split('{')[0].strip()
            # Remove trailing comma if present (e.g. from "Dodge, Sr.")
            clean_name = clean_name.rstrip(",.")

            # 1. Full Name
            name_index[clean_name].append(pid)

            # 2. Base Name (remove suffix)
            # Handle "Jr", "Sr", "III", "IV", "Esq."
            # Using regex to remove suffix at the end
            base_name = re.sub(r',?\s+(Jr\.?|Sr\.?|III|IV|Esq\.?)$', '', clean_name, flags=re.IGNORECASE)
            if base_name != clean_name:
                name_index[base_name].append(pid)

            # 3. Variations on Base Name
            # "William Earl Dodge" -> "William Dodge"
            # "William E. Dodge" -> "William Dodge"
            # "William E. Dodge" -> "William E. Dodge" (Already covered by base_name if no suffix)

            parts = base_name.split()

            # If name has middle parts (more than 2 words)
            if len(parts) > 2:
                # First Last
                short_name = f"{parts[0]} {parts[-1]}"
                name_index[short_name].append(pid)

                # If middle initial is used in text "William E. Dodge", we want that to match
                # "William Earl Dodge" from database.
                # So if DB has "William Earl Dodge", we can index "William E. Dodge" too?
                # No, we index the full name. We rely on text scanning to pick up "William E. Dodge"
                # and then we need to match it.
                # If DB is "William Earl Dodge", and text says "William E. Dodge",
                # valid_candidates will have "William E. Dodge".
                # We need "William E. Dodge" in known_names to trigger a match?
                # Yes.

                # Generate Middle Initial Variant
                # "William Earl Dodge" -> "William E. Dodge"
                middle_part = parts[1]
                if len(middle_part) > 1 and middle_part[0].isalpha():
                    initial_name = f"{parts[0]} {middle_part[0]}. {parts[-1]}"
                    name_index[initial_name].append(pid)

        return name_index

    def _scan_text_for_mentions(self, text, name_index, source_profile):
        """
        Scans a text block for mentions of names in the index.
        Returns a list of related_link objects.
        """
        if not text:
            return []

        links = []
        # known_names = set(name_index.keys())

        # Keywords for relationship types
        # Added: Uncle, Aunt, Nephew, Niece, Executor, Witness
        keywords = {
            "partner": "Business Partner",
            "business": "Business Partner",
            "firm": "Business Partner",
            "colleague": "Business Partner",
            "married": "Spouse",
            "wife": "Spouse",
            "husband": "Spouse",
            "spouse": "Spouse",
            "wed": "Spouse",
            "cousin": "Cousin",
            "friend": "Friend",
            "neighbor": "Neighbor",
            "associate": "Associate",
            "classmate": "Classmate",
            "tutor": "Tutor",
            "student": "Student",
            "enemy": "Rival",
            "rival": "Rival",
            "uncle": "Relative",
            "aunt": "Relative",
            "nephew": "Relative",
            "niece": "Relative",
            "executor": "Legal Associate",
            "witness": "Legal Associate",
            "legacy": "Relative",
            "mother-in-law": "In-Law",
            "father-in-law": "In-Law",
            "son-in-law": "In-Law",
            "daughter-in-law": "In-Law",
            "brother-in-law": "In-Law",
            "sister-in-law": "In-Law",
            "step-mother": "Step-Parent",
            "step-father": "Step-Parent",
            "step-son": "Step-Child",
            "step-daughter": "Step-Child",
            "fiancé": "Fiancé",
            "fiancee": "Fiancé",
            "betrothed": "Fiancé",
            "mentor": "Professional",
            "apprentice": "Professional",
            "godfather": "Godparent",
            "godmother": "Godparent",
            "godson": "Godparent",
            "goddaughter": "Godparent"
        }

        # Improved Candidate Extraction
        # Pattern: Capitalized Words sequence
        # \b(?:[A-Z]\.?|[A-Z][a-z]+)(?:\s+(?:[A-Z]\.?|[A-Z][a-z]+))+\b
        name_pattern = r'\b(?:[A-Z]\.?|[A-Z][a-z]+)(?:\s+(?:[A-Z]\.?|[A-Z][a-z]+))+\b'
        candidates = set(re.findall(name_pattern, text))

        # Filter candidates that are known names
        # Use name_index.keys() which contains all variations
        valid_candidates = candidates.intersection(set(name_index.keys()))

        if not valid_candidates:
            return []

        # Verification and Context
        clauses = re.split(r'[.;,]', text)
        source_born = self._get_birth_year(source_profile)
        found_ids = set()

        id_map = {p['id']: p for p in self.family_data}

        for clause in clauses:
            if not clause.strip(): continue

            for name in valid_candidates:
                if name not in clause: continue # Fast string check

                # Exact word boundary check
                if not re.search(r'\b' + re.escape(name) + r'\b', clause):
                    continue

                # Ambiguity Resolution Strategy
                # 1. Get all potential IDs for this name
                potential_ids = name_index[name]

                # 2. Filter out self
                potential_ids = [pid for pid in potential_ids if pid != source_profile['id']]

                if not potential_ids:
                    continue

                # 3. Decision Logic
                target_id = None

                if len(potential_ids) == 1:
                    # Unambiguous (only 1 candidate) - Accept it regardless of date (could be ancestor)
                    target_id = potential_ids[0]
                else:
                    # Ambiguous (multiple candidates) - Apply Strict Date Filter
                    candidates_in_range = []
                    for pid in potential_ids:
                        target_p = id_map.get(pid)
                        if not target_p: continue

                        target_born = self._get_birth_year(target_p)

                        # Strict check for disambiguation (must be within 60 years)
                        if source_born and target_born:
                            diff = abs(source_born - target_born)
                            if diff <= 60:
                                candidates_in_range.append(pid)
                        else:
                            # If dates unknown, we can't safely disambiguate by date
                            # Treat as ambiguous unless it's the only one?
                            # Safe bet: skip if date unknown and ambiguous
                            pass

                    if len(candidates_in_range) == 1:
                        target_id = candidates_in_range[0]
                    else:
                         # Priority: Real Profile over Child Entry
                        real_candidates = [pid for pid in candidates_in_range if "_c" not in pid]
                        if len(real_candidates) == 1:
                             target_id = real_candidates[0]
                        else:
                            # Still ambiguous
                            self.ariadne_log["ambiguous"][name].update(potential_ids)
                            continue

                if target_id in found_ids: continue

                found_ids.add(target_id)

                # 4. Contemporary Check for Relation Type (Loose Check)
                # Now that we have the target, we check if they are "contemporary" for the purpose
                # of inferring "Friend", "Partner", etc. vs just "Mentioned".
                target_p = id_map.get(target_id)
                target_born = self._get_birth_year(target_p)
                is_contemporary = True
                if source_born and target_born:
                     if abs(source_born - target_born) > 80: # Keep loose check for 'Mentioned' fallback
                         is_contemporary = False

                # Determine Type
                rel_type = "Mentioned"
                lower_clause = clause.lower()

                for k, v in keywords.items():
                    if re.search(r'\b' + re.escape(k) + r'\b', lower_clause):
                        # Special handling for Spousal/Partner keywords vs Non-Contemporary matches
                        if v in ["Spouse", "Business Partner", "Friend", "Classmate"] and not is_contemporary:
                            rel_type = "Mentioned"
                        else:
                            rel_type = v
                        break

                # Context sentence (find full sentence containing the clause)
                full_sentences = self.split_sentences(text)
                source_sentence = clause.strip()
                for s in full_sentences:
                    if clause.strip() in s:
                        source_sentence = s
                        break

                links.append({
                    "target_id": target_id,
                    "relation_type": rel_type,
                    "source_text": source_sentence.strip()
                })

                self.ariadne_log["new_links"] += 1
                self.ariadne_log["clusters"][name] += 1

        return links

    def _find_mentions(self):
        print("--- Ariadne: Weaving Connections ---")

        # Build Name Index
        name_index = self._build_name_index()

        # Track discoveries for the journal
        self.ariadne_log = {
            "ambiguous": defaultdict(set),
            "clusters": defaultdict(int),
            "new_links": 0
        }

        count = 0

        for p in self.family_data:
            p['related_links'] = []
            notes = p['story']['notes']
            if not notes: continue

            # Pass the full name_index (with ambiguities)
            links = self._scan_text_for_mentions(notes, name_index, p)
            p['related_links'].extend(links)
            count += len(links)

        print(f"Ariadne found {count} text-based connections.")

        # ---------------------------
        # Post-Processing: Symmetry
        # ---------------------------
        print("--- Ariadne: Stitching Reverse Links ---")
        reverse_count = 0
        id_map = {p['id']: p for p in self.family_data}

        for p in self.family_data:
            source_id = p['id']
            source_name = p['name']

            # Iterate over a copy of links
            for link in list(p['related_links']):
                target_id = link['target_id']
                rel_type = link['relation_type']
                source_text = link['source_text']

                target_p = id_map.get(target_id)
                if not target_p: continue

                if 'related_links' not in target_p:
                    target_p['related_links'] = []

                # Check if reverse link already exists (to prevent duplicates)
                exists = False
                for r_link in target_p['related_links']:
                    if r_link['target_id'] == source_id:
                        # Check context - if it's already linked, maybe we don't need another?
                        # Or checking source_text match might be too specific if text differs?
                        # Let's say if A linked to B, B should link to A.
                        # If B already links to A, we skip.
                        exists = True
                        break

                if not exists:
                    # Determine reverse type
                    rev_type = rel_type

                    if rel_type == "Mentioned":
                        rev_type = "Mentioned by"
                    elif rel_type == "Parent":
                        rev_type = "Child"
                    elif rel_type == "Child":
                        rev_type = "Parent"
                    elif rel_type == "In-Law":
                        rev_type = "In-Law" # Could be specific, but In-Law is safe generic
                    elif rel_type == "Step-Parent":
                        rev_type = "Step-Child"
                    elif rel_type == "Step-Child":
                        rev_type = "Step-Parent"
                    elif rel_type == "Godparent":
                        rev_type = "Godchild"

                    target_p['related_links'].append({
                        "target_id": source_id,
                        "relation_type": rev_type,
                        "source_text": f"Mentioned in {source_name}'s notes: \"{source_text[:50]}...\""
                    })
                    reverse_count += 1

        print(f"Ariadne added {reverse_count} reverse connections.")

        # Log Interesting Findings (Console for now)
        print("\n--- Ariadne's Notebook ---")
        if self.ariadne_log["ambiguous"]:
            print(f"Ambiguous Names Skipped: {len(self.ariadne_log['ambiguous'])}")
            # Print top 5 ambiguous
            for k in list(self.ariadne_log["ambiguous"].keys())[:5]:
                print(f"  - {k}: {self.ariadne_log['ambiguous'][k]}")

        # Check for potential clusters (frequently mentioned names)
        sorted_clusters = sorted(self.ariadne_log["clusters"].items(), key=lambda x: x[1], reverse=True)
        print("Top Mentioned People:")
        for name, freq in sorted_clusters[:5]:
            print(f"  - {name}: {freq} mentions")

    def _get_country_from_location(self, location):
        if not location or location == "Unknown":
            return "Unknown"

        loc_lower = location.lower()

        # UK/Europe
        if any(x in loc_lower for x in ["england", "uk", "britain", "london", "liverpool", "plymouth", "bristol", "southampton", "dorset", "essex", "suffolk", "kent", "somerset", "lincolnshire", "yarmouth", "cowes", "gravesend"]):
            return "UK"
        if "france" in loc_lower: return "France"
        if "germany" in loc_lower or "prussia" in loc_lower: return "Germany"
        if "holland" in loc_lower or "netherlands" in loc_lower or "amsterdam" in loc_lower or "leyden" in loc_lower or "leiden" in loc_lower: return "Netherlands"

        # USA/America
        # Keywords
        usa_keywords = [
            "america", "usa", "united states", "new england",
            "massachusetts", "mass", "connecticut", "new york", "virginia", "pennsylvania", "new jersey", "maryland", "carolina", "georgia", "vermont", "new hampshire", "rhode island", "maine", "ohio", "illinois", "michigan",
            "boston", "salem", "hartford", "jamestown", "plymouth", "new haven", "windsor", "wethersfield", "watertown", "dorchester", "roxbury"
        ]
        if any(x in loc_lower for x in usa_keywords):
            return "USA"

        # State Abbreviations (with boundary checks to avoid false positives like 'me', 'pa')
        # We look for " CT", ",CT", " CT " or end of string.
        # Simplest is regex for \b(CT|MA|NY|NJ|PA|VA|RI|VT|NH|ME|DE|MD|SC|NC|GA|OH|IL|MI)\b
        if re.search(r'\b(ct|ma|ny|nj|pa|va|ri|vt|nh|me|de|md|sc|nc|ga|oh|il|mi)\b', loc_lower):
            return "USA"

        return "Unknown"

    def _generate_voyage_context(self, voyage, profile):
        # Determine effective year for context
        year = None

        # 1. Try explicit voyage year
        if voyage.get("year") and voyage["year"] != "Unknown":
            try:
                year = int(voyage["year"])
            except:
                pass

        # 2. Try Ship Year Built (as approximation of era)
        if not year and voyage.get("specs") and voyage["specs"].get("year_built"):
             try:
                 # "1630" or "c. 1630"
                 yb_str = str(voyage["specs"]["year_built"])
                 match = re.search(r'\d{4}', yb_str)
                 if match:
                     year = int(match.group(0)) + 5 # Assume voyage is slightly after build
             except:
                 pass

        # 3. Try Ancestor Migration Logic (Birth/Death)
        is_estimated = False
        if not year:
            b_year = profile["vital_stats"].get("born_year_int")
            d_year = profile["vital_stats"].get("died_year_int")
            if b_year:
                # Assume migration happened in early adulthood or childhood?
                # If we don't know, Mid-Life is a safe statistical guess for "Era" context
                if d_year:
                    year = int((b_year + d_year) / 2)
                else:
                    year = b_year + 30
                is_estimated = True

        if not year:
            return None

        # Determine Route
        dep_country = self._get_country_from_location(voyage.get("departure", ""))
        if dep_country == "Unknown":
             dep_country = self._get_country_from_location(profile["vital_stats"].get("born_location", ""))

        arr_country = self._get_country_from_location(voyage.get("arrival", ""))
        if arr_country == "Unknown":
             arr_country = self._get_country_from_location(profile["vital_stats"].get("died_location", ""))

        # Context Logic
        context = {
            "year_used": year,
            "is_estimated": is_estimated,
            "duration": "Unknown",
            "conditions": "Unknown",
            "era_label": "Unknown"
        }

        # Trans-Atlantic (UK/Europe -> USA)
        if (dep_country in ["UK", "Netherlands", "Germany", "France"]) and (arr_country == "USA"):
            if year < 1650: # Great Migration Era
                context["duration"] = "8-12 weeks"
                context["conditions"] = "Extremely cramped and hazardous. High risk of scurvy and infectious disease. Passengers slept in 'tween decks' with little ventilation."
                context["era_label"] = "The Great Migration"
            elif year < 1700:
                context["duration"] = "8-10 weeks"
                context["conditions"] = "Slightly larger vessels but still dangerous. Poor sanitation and limited rations were standard."
                context["era_label"] = "Colonial Era"
            elif year < 1800:
                context["duration"] = "6-8 weeks"
                context["conditions"] = "Regular packet ships began to operate. Conditions improved slightly but storms remained a major threat."
                context["era_label"] = "18th Century Passage"
            elif year < 1850:
                context["duration"] = "4-6 weeks"
                context["conditions"] = "The era of fast clipper ships and packets. Steerage conditions remained poor/crowded for immigrants."
                context["era_label"] = "Age of Sail"
            else:
                context["duration"] = "10-14 days"
                context["conditions"] = "Early steamships offered consistent travel times. Steerage was still crowded, but the journey was significantly shorter and safer."
                context["era_label"] = "Steamship Era"

        return context

    def clean_and_save(self):
        # 1. First pass: separate "Real" profiles from "Child" entries
        real_profiles = {}
        child_entries = []

        for p in self.family_data:
            if p.get("metadata", {}).get("is_child_entry"):
                child_entries.append(p)
            else:
                pid = p['id']
                if pid not in real_profiles:
                    real_profiles[pid] = p

        # 2. Process Child Entries:
        # If a child entry matches a Real Profile by Name, discard the child entry (the real one is better).
        # Otherwise, keep the child entry.

        # Build Name Map for Real Profiles
        real_name_map = {}
        for pid, p in real_profiles.items():
            # Normalize name: remove parens, extra spaces
            norm = re.sub(r'\(.*?\)', '', p['name']).strip().lower()
            real_name_map[norm] = pid

        final_profiles = list(real_profiles.values())

        for child in child_entries:
            # Clean child name for matching
            # Child name might be "Grace Dodge (John) Olmsted..." -> "Grace Dodge Olmsted"
            # It's tricky.
            # Simple check: if "Grace Dodge" is in the real profile name?

            # Simple normalization
            c_name = re.sub(r'\[.*?\]', '', child['name']) # remove spouses in brackets
            c_name = re.sub(r'\(.*?\)', '', c_name)
            # Remove suffixes like ", 1st son", ", 2nd daughter", etc.
            c_name = re.sub(r',\s*\d+(?:st|nd|rd|th)?\s+(?:son|daughter|child)', '', c_name, flags=re.IGNORECASE)
            c_name = re.sub(r',\s+(?:son|daughter|child)', '', c_name, flags=re.IGNORECASE)
            c_name = c_name.strip().lower()

            match_found = False
            for real_name, real_id in real_name_map.items():
                # Check for strong match
                if c_name == real_name or (len(c_name) > 5 and c_name in real_name):
                    # Link parent to this REAL profile instead of the child entry
                    parent_id = child["metadata"]["parent_id"]

                    # Add parent/child link
                    if parent_id in real_profiles:
                         parent = real_profiles[parent_id]
                         if 'relations' not in parent: parent['relations'] = {"children": [], "parents": [], "spouses": []}
                         if real_id not in parent['relations']['children']:
                             parent['relations']['children'].append(real_id)

                    real_p = real_profiles[real_id]
                    if 'relations' not in real_p: real_p['relations'] = {"children": [], "parents": [], "spouses": []}

                    # Prevent linking self as parent (e.g. if child name matches parent name)
                    if parent_id != real_id:
                        if parent_id not in real_p['relations']['parents']:
                            real_p['relations']['parents'].append(parent_id)

                    match_found = True
                    break

            if not match_found:
                # Keep this child entry as a new profile
                final_profiles.append(child)

                # Link to parent
                parent_id = child["metadata"]["parent_id"]
                if parent_id in real_profiles:
                     parent = real_profiles[parent_id]
                     if 'relations' not in parent: parent['relations'] = {"children": [], "parents": [], "spouses": []}
                     # Add this child ID
                     if child['id'] not in parent['relations']['children']:
                         parent['relations']['children'].append(child['id'])

                # Link child to parent
                if 'relations' not in child: child['relations'] = {"children": [], "parents": [], "spouses": []}
                child['relations']['parents'].append(parent_id)

        self.family_data = final_profiles

        # Re-deduplicate just in case?
        # self.family_data already unique by logic above.

        self.link_family_members()
        self._analyze_naming_patterns()
        self._find_mentions()

        # Initialize Services
        geocoder = GeocodingService()
        ship_service = ShipEnrichmentService()

        print("--- Fetching Hero Images, Geocoding, and Ship Info ---")

        # Calculate Associate Social Capital (Frequency)
        associate_counts = defaultdict(int)
        for p in self.family_data:
            if "associates" in p["story"]:
                for assoc in p["story"]["associates"]:
                    # Normalize name for counting (simple case-insensitive)
                    associate_counts[assoc["name"].lower()] += 1

        # --- Shipmates Logic ---
        ship_manifest = defaultdict(list)
        for p in self.family_data:
            if "voyages" in p["story"]:
                for voyage in p["story"]["voyages"]:
                    s_name = voyage.get("ship_name")
                    if s_name and s_name != "Unknown":
                        ship_manifest[s_name].append({
                            "id": p["id"],
                            "name": p["name"]
                        })

        final_list = []
        
        for p in self.family_data:
            # Inject Associate Frequency
            if "associates" in p["story"]:
                for assoc in p["story"]["associates"]:
                    assoc["count"] = associate_counts.get(assoc["name"].lower(), 1)

            # Geocode
            born_loc = p["vital_stats"]["born_location"]
            died_loc = p["vital_stats"]["died_location"]

            p["vital_stats"]["born_coords"] = geocoder.geocode(born_loc)
            p["vital_stats"]["died_coords"] = geocoder.geocode(died_loc)

            # Geocode Life Events
            for event in p["story"]["life_events"]:
                if "location" in event and event["location"] != "Unknown":
                    event["coords"] = geocoder.geocode(event["location"])

            # Enrich Voyage Data
            if "voyages" in p["story"]:
                for voyage in p["story"]["voyages"]:
                    # 1. Generalize Route (add country fields)
                    voyage["departure_country"] = self._get_country_from_location(voyage.get("departure"))
                    if voyage["departure_country"] == "Unknown":
                        voyage["departure_country"] = self._get_country_from_location(p["vital_stats"].get("born_location"))

                    voyage["arrival_country"] = self._get_country_from_location(voyage.get("arrival"))
                    if voyage["arrival_country"] == "Unknown":
                        voyage["arrival_country"] = self._get_country_from_location(p["vital_stats"].get("died_location"))

                    # 2. Enrich Ship Specs
                    ship_name = voyage.get("ship_name")
                    if ship_name and ship_name != "Unknown":
                        specs = ship_service.enrich_ship(ship_name)
                        if specs:
                            voyage["specs"] = specs

                        # Add Shipmates
                        if ship_name in ship_manifest:
                            mates = [m for m in ship_manifest[ship_name] if m["id"] != p["id"]]
                            if mates:
                                voyage["shipmates"] = mates

                    # 3. Generate Historical Context (Visionary Layer)
                    voyage["context"] = self._generate_voyage_context(voyage, p)

            # Fetch Image
            born_year = p["vital_stats"].get("born_year_int")
            hero_image = self.fetch_wikimedia_image(born_loc, born_year)

            # Extract Tags
            tags = self.extract_tags(p)

            final_profile = {
                "id": p["id"],
                "name": p["name"],
                "lineage": p.get("lineage", "Unknown"),
                "generation": p["generation"],
                "vital_stats": p["vital_stats"],
                "story": {
                    "notes": p["story"]["notes"],
                    "voyages": p["story"].get("voyages", []),
                    "life_events": p["story"].get("life_events", []),
                    "tags": tags,
                    "associates": p["story"].get("associates", []),
                    "naming_echo": p["story"].get("naming_echo")
                },
                "hero_image": hero_image,
                "relations": p.get("relations", {}),
                "related_links": p.get("related_links", []),
                "metadata": {
                    "source_ref": p["metadata"]["source_id"],
                    "location_in_doc": f"Paragraph #{p['metadata']['doc_paragraph_index']}"
                }
            }
            final_list.append(final_profile)

        # Save Cache
        self.save_cache()
        geocoder.save_cache()
        ship_service.save_cache()

        output_filename = "kinship-app/src/family_data.json"
        with open(output_filename, "w", encoding='utf-8') as f:
            json.dump(final_list, f, indent=4, ensure_ascii=True)
        
        print(f"Data saved to {output_filename}")

if __name__ == "__main__":
    files = {
        "Paternal": "GENEALOGY DSD Paternal Ancestry.docx",
        "Maternal": "GENEALOGY DSD Maternal Ancestry.docx"
    }

    pipeline = GenealogyTextPipeline()

    for lineage, filename in files.items():
        if os.path.exists(filename):
            pipeline.parse_document(filename, lineage)
        else:
            print(f"Error: Could not find {filename}")

    pipeline.clean_and_save()
